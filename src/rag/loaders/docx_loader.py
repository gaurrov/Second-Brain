"""
DOCX text extraction using python-docx.

DOCX has no native "page" concept (pagination is a rendering-time
concern), so the whole document is returned as a single LoadedPage with
page_number=None. Paragraphs are joined with newlines to preserve
paragraph boundaries for the text splitter. Tables are extracted
row-by-row with pipe separators for structured data.
"""
import logging
from pathlib import Path

import docx
from docx.opc.exceptions import PackageNotFoundError

from src.core.exceptions import TextExtractionException
from src.rag.loaders.base_loader import LoadedPage

logger = logging.getLogger("second_brain.loaders.docx")


class DOCXLoader:
    def load(self, file_path: Path) -> list[LoadedPage]:
        if not file_path.exists():
            raise TextExtractionException(f"DOCX file not found: {file_path}")
        if not file_path.is_file():
            raise TextExtractionException(f"Path is not a file: {file_path}")

        try:
            document = docx.Document(str(file_path))
        except (PackageNotFoundError, OSError, KeyError) as exc:
            raise TextExtractionException(f"Could not open DOCX file: {exc}") from exc

        paragraphs: list[str] = []
        skipped = 0

        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
            elif para.text:
                skipped += 1

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                row_text = " | ".join(cells)
                if row_text:
                    paragraphs.append(row_text)

        if skipped > 0:
            logger.debug("DOCX extraction: skipped %d empty paragraphs", skipped)

        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            raise TextExtractionException("No extractable text found in this DOCX file.")

        return [LoadedPage(page_number=None, content=full_text)]
